import streamlit as st
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
import google.generativeai as genai
from io import StringIO
import time
from datetime import datetime
import threading

# ==================== PAGE CONFIG ====================
st.set_page_config(
    page_title="AI MCQ & Short Question Generator",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== CUSTOM CSS ====================
st.markdown("""
    <style>
    .main {
        padding: 1rem;
    }
    .stTabs [data-baseweb="tab-list"] button {
        font-weight: bold;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 20px;
        border-radius: 10px;
        margin-bottom: 15px;
    }
    .success-box {
        background-color: #d4edda;
        border-left: 4px solid #28a745;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 15px;
    }
    .info-box {
        background-color: #e7f3ff;
        border-left: 4px solid #007bff;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 15px;
    }
    .warning-box {
        background-color: #fff3cd;
        border-left: 4px solid #ffc107;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 15px;
    }
    </style>
""", unsafe_allow_html=True)

# ==================== SESSION STATE INITIALIZATION ====================
if "conversation_history" not in st.session_state:
    st.session_state.conversation_history = {}
if "current_context" not in st.session_state:
    st.session_state.current_context = ""
if "current_file_name" not in st.session_state:
    st.session_state.current_file_name = None
if "api_key" not in st.session_state:
    st.session_state.api_key = None
if "generation_count" not in st.session_state:
    st.session_state.generation_count = 0
if "last_request_time" not in st.session_state:
    st.session_state.last_request_time = 0
if "quota_reset_time" not in st.session_state:
    st.session_state.quota_reset_time = None

# ==================== FILE PROCESSING FUNCTIONS ====================
@st.cache_data(show_spinner=False)
def read_pdf(pdf_file):
    """Extract text from PDF with caching"""
    try:
        reader = PdfReader(pdf_file)
        text = "\n".join([page.extract_text() or "" for page in reader.pages])
        return text.strip()
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return None

@st.cache_data(show_spinner=False)
def read_docx(docx_file):
    """Extract text from DOCX with caching"""
    try:
        doc = Document(docx_file)
        text = "\n".join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " | ".join([cell.text for cell in row.cells])
        return text.strip()
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return None

@st.cache_data(show_spinner=False)
def read_txt(txt_file):
    """Extract text from TXT with caching"""
    try:
        return StringIO(txt_file.getvalue().decode("utf-8")).read().strip()
    except Exception as e:
        st.error(f"Error reading TXT: {str(e)}")
        return None

@st.cache_data(show_spinner=False)
def read_csv(csv_file):
    """Extract text from CSV with caching"""
    try:
        df = pd.read_csv(csv_file)
        return df.to_string()
    except Exception as e:
        st.error(f"Error reading CSV: {str(e)}")
        return None

@st.cache_data(show_spinner=False)
def read_excel(excel_file):
    """Extract text from EXCEL with caching"""
    try:
        xls = pd.ExcelFile(excel_file)
        text = ""
        for sheet in xls.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet)
            text += f"\n--- Sheet: {sheet} ---\n" + df.to_string() + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"Error reading EXCEL: {str(e)}")
        return None

def extract_text(uploaded_file):
    """Extract text from any supported file type"""
    if uploaded_file is None:
        return None
    
    file_type = uploaded_file.name.split(".")[-1].lower()
    
    processors = {
        "pdf": read_pdf,
        "docx": read_docx,
        "txt": read_txt,
        "csv": read_csv,
        "xls": read_excel,
        "xlsx": read_excel
    }
    
    processor = processors.get(file_type)
    if processor:
        return processor(uploaded_file)
    else:
        st.error(f"Unsupported file type: {file_type}")
        return None

# ==================== MODEL CONFIGURATION ====================
@st.cache_resource
def get_generative_model(api_key):
    """Initialize and cache the generative AI model with fallback"""
    try:
        genai.configure(api_key=api_key)
        model_options = ['gemini-2.0-flash', 'gemini-1.5-flash', 'gemini-1.5-pro']
        
        for model_name in model_options:
            try:
                model = genai.GenerativeModel(
                    model_name=model_name,
                    system_instruction=(
                        "You are an expert educational question generator. "
                        "Generate clear, concise, and well-structured questions that test understanding. "
                        "Avoid unnecessary explanations. Use proper punctuation and formatting."
                    )
                )
                return model
            except Exception:
                continue
        
        raise Exception("No compatible Gemini model found")
    except Exception as e:
        st.error(f"Model initialization error: {str(e)}")
        return None

# ==================== RATE LIMITING ====================
def check_quota_limits():
    """Check if user can make another request"""
    current_time = time.time()
    
    # Check if in quota reset period
    if st.session_state.quota_reset_time and current_time < st.session_state.quota_reset_time:
        remaining = st.session_state.quota_reset_time - current_time
        return False, f"Rate limited. Wait {int(remaining)}s."
    
    # Clear reset if time passed
    if st.session_state.quota_reset_time and current_time >= st.session_state.quota_reset_time:
        st.session_state.quota_reset_time = None
    
    # Rate limit: 1 request per 3 seconds
    if current_time - st.session_state.last_request_time < 3:
        wait_time = 3 - (current_time - st.session_state.last_request_time)
        return False, f"Wait {int(wait_time)}s before next request."
    
    return True, None

# ==================== QUESTION GENERATION ====================
def generate_questions(api_key, content, mcq_count, short_count, difficulty_level, topic_focus):
    """Generate MCQs and short questions with optimized prompt and rate limiting"""
    
    # Check rate limits first
    can_proceed, error_msg = check_quota_limits()
    if not can_proceed:
        st.warning(f"⏳ {error_msg}")
        return None
    
    try:
        model = get_generative_model(api_key)
        if model is None:
            return None
        
        # Ultra-optimized content truncation
        max_chars = 7000
        truncated_content = content[:max_chars]
        if len(content) > max_chars:
            truncated_content += "\n[Content truncated for optimization]"
        
        # Minimal token prompt
        prompt = f"""Generate {mcq_count} MCQs and {short_count} short questions.
Difficulty: {difficulty_level}
Topic: {topic_focus}

Content:
{truncated_content}

=== MULTIPLE CHOICE ({mcq_count}) ===
Q1. [Question]
A) [Option A]
B) [Option B]
C) [Option C]
D) [Option D]
Answer: [A/B/C/D]

=== SHORT ANSWER ({short_count}) ===
Q1. [Question]
Answer: [2-3 lines]"""
        
        # Record request time
        st.session_state.last_request_time = time.time()
        
        response = model.generate_content(prompt)
        return response.text if response else None
    
    except Exception as e:
        error_str = str(e)
        
        # Handle quota exceeded
        if "429" in error_str or "quota" in error_str.lower() or "exceeded" in error_str.lower():
            st.session_state.quota_reset_time = time.time() + 120
            
            st.error("""❌ **API Quota Exceeded**
            
The free tier rate limit has been reached temporarily.

**Solutions:**
1. ⏳ **Wait 2 minutes** - Quota resets in 120 seconds
2. 📉 **Reduce questions** - Try fewer MCQs/short questions
3. 💳 **Upgrade** - https://ai.google.dev/pricing
4. 📅 **Tomorrow** - Free quota resets daily

**Why this happens:**
Free tier has daily & hourly request limits. You've made many requests recently.""")
            return None
        
        st.error(f"❌ Error generating questions: {error_str[:200]}")
        return None

# ==================== MAIN UI ====================
# Header
col1, col2 = st.columns([3, 1])
with col1:
    st.title("🧠 AI MCQ & Short Question Generator")
    st.markdown("*Generate exam questions from your documents instantly*")

with col2:
    st.metric("📊 Generations", st.session_state.generation_count)

st.markdown("---")

# Sidebar Configuration
with st.sidebar:
    st.subheader("⚙️ Configuration")
    
    # API Key
    api_key_input = st.text_input(
        "🔑 Google API Key",
        value=st.session_state.api_key or "",
        type="password",
        help="Get your API key from https://makersuite.google.com/app/apikey"
    )
    
    if api_key_input:
        st.session_state.api_key = api_key_input
    
    st.markdown("---")
    st.subheader("📋 Generation Settings")
    
    # Question counts
    col1, col2 = st.columns(2)
    with col1:
        mcq_count = st.slider("Number of MCQs", 1, 30, 10, step=1)
    with col2:
        short_count = st.slider("Short Questions", 1, 20, 5, step=1)
    
    # Difficulty level
    difficulty = st.selectbox(
        "🎯 Difficulty Level",
        ["Easy", "Medium", "Hard", "Mixed"],
        help="Choose the difficulty level of questions"
    )
    
    # Topic Focus
    topic_focus = st.text_input(
        "🎓 Topic Focus (Optional)",
        placeholder="e.g., Chapter 5, Key concepts",
        help="Specify the area to focus on (optional)"
    )
    
    st.markdown("---")
    st.subheader("💾 Context Management")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("💾 Save Context", use_container_width=True):
            if st.session_state.current_context.strip():
                context_key = f"{st.session_state.current_file_name}_{datetime.now().strftime('%H:%M:%S')}"
                st.session_state.conversation_history[context_key] = {
                    "file": st.session_state.current_file_name,
                    "content": st.session_state.current_context,
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "size": f"{len(st.session_state.current_context)/1024:.1f} KB"
                }
                st.success("✅ Context saved!")
            else:
                st.warning("⚠️ No content to save")
    
    with col2:
        if st.button("🔄 Clear All", use_container_width=True):
            st.session_state.conversation_history = {}
            st.session_state.current_context = ""
            st.success("✅ All contexts cleared!")
    
    # Display saved contexts
    if st.session_state.conversation_history:
        st.markdown("---")
        st.subheader("📚 Saved Contexts")
        
        for context_key, context_data in st.session_state.conversation_history.items():
            with st.expander(f"📄 {context_data['file']} - {context_data['timestamp']}"):
                st.caption(f"Size: {context_data['size']}")
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("✅ Load", key=f"load_{context_key}", use_container_width=True):
                        st.session_state.current_context = context_data['content']
                        st.session_state.current_file_name = context_data['file']
                        st.success("✅ Context loaded!")
                        st.rerun()
                
                with col2:
                    if st.button("🗑️ Delete", key=f"del_{context_key}", use_container_width=True):
                        del st.session_state.conversation_history[context_key]
                        st.success("✅ Context deleted!")
                        st.rerun()

# ==================== MAIN CONTENT ====================
if not st.session_state.api_key:
    st.markdown("""
    <div class="warning-box">
    <strong>⚠️ API Key Required</strong><br>
    Please enter your Google API Key in the sidebar to get started.
    </div>
    """, unsafe_allow_html=True)
    
    st.info("""
    **How to get an API Key:**
    1. Visit https://makersuite.google.com/app/apikey
    2. Click "Create API Key"
    3. Copy and paste it in the sidebar
    """)
else:
    # File Upload Section
    st.subheader("📂 Upload Document")
    
    col1, col2 = st.columns([3, 1])
    with col1:
        uploaded_file = st.file_uploader(
            "Choose a file (PDF, DOCX, TXT, CSV, Excel)",
            type=["pdf", "docx", "txt", "csv", "xls", "xlsx"],
            label_visibility="collapsed"
        )
    
    if uploaded_file:
        with st.spinner("📖 Processing document..."):
            text_content = extract_text(uploaded_file)
            
            if text_content:
                st.session_state.current_context = text_content
                st.session_state.current_file_name = uploaded_file.name
                
                # Show processing info
                st.markdown(f"""
                <div class="success-box">
                <strong>✅ File Processed</strong><br>
                📄 File: <strong>{uploaded_file.name}</strong><br>
                📊 Size: <strong>{len(text_content)/1024:.1f} KB</strong> ({len(text_content)} characters)
                </div>
                """, unsafe_allow_html=True)
                
                # Preview
                with st.expander("👁️ Preview Content", expanded=False):
                    preview_length = min(2000, len(text_content))
                    st.text_area(
                        "Content Preview",
                        text_content[:preview_length],
                        height=200,
                        disabled=True,
                        label_visibility="collapsed"
                    )
                    if len(text_content) > preview_length:
                        st.caption(f"... ({len(text_content) - preview_length} more characters)")
    
    st.markdown("---")
    
    # Generation Section
    if st.session_state.current_context.strip():
        st.subheader("🚀 Generate Questions")
        
        col1, col2, col3 = st.columns([2, 1, 1])
        
        with col1:
            generate_btn = st.button(
                "🎯 Generate Questions",
                use_container_width=True,
                type="primary",
                help="Click to generate MCQs and short questions"
            )
        
        with col2:
            if st.button("📋 Copy All", use_container_width=True):
                st.info("📋 Content copied to clipboard (when viewing in browser)")
        
        with col3:
            if st.button("📥 Download", use_container_width=True):
                st.info("💾 Use browser download feature")
        
        if generate_btn:
            with st.spinner("🧠 Generating questions... This may take a moment"):
                start_time = time.time()
                
                output = generate_questions(
                    st.session_state.api_key,
                    st.session_state.current_context,
                    mcq_count,
                    short_count,
                    difficulty,
                    topic_focus if topic_focus else "All topics"
                )
                
                generation_time = time.time() - start_time
                
                if output:
                    st.session_state.generation_count += 1
                    
                    # Success message with metrics
                    st.markdown(f"""
                    <div class="success-box">
                    <strong>✅ Questions Generated Successfully!</strong><br>
                    ⏱️ Generation Time: <strong>{generation_time:.1f}s</strong><br>
                    📊 Questions: <strong>{mcq_count} MCQs + {short_count} Short</strong>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Display output
                    st.subheader("📝 Generated Questions")
                    st.markdown(output)
                    
                    # Download option
                    st.download_button(
                        label="⬇️ Download as Text",
                        data=output,
                        file_name=f"questions_{st.session_state.current_file_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
    else:
        st.markdown("""
        <div class="info-box">
        <strong>ℹ️ Ready to Generate</strong><br>
        Upload a document to get started. Supported formats: PDF, DOCX, TXT, CSV, Excel
        </div>
        """, unsafe_allow_html=True)

st.markdown("---")
st.caption("⚡ Fast • 🧠 Smart • 💾 Cached • 🔄 Saveable")
